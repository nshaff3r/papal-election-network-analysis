import os
from docx import Document
import re

class Character:
    def __init__(self, name):
        self.name = name
        self.other = ""
        self.relations = {}
        self.relationsFiltered = {}
people = [
    "charles viii, king of france, of the house of anjou",
    "giambattista orsini, cardinal deacon of s. giorgio in velabro",
    "father mario maffei",
    "isabella i, queen of castile and léon, queen consort of aragon, naples, sicily and valencia, of the house of trastámara",
    "domenico della rovere di vinovo, cardinal of san clemente",
    "papal chaplain alessandro farnese",
    "federico (da montefeltro) sanseverino,",
    "cesare “valentino” borgia, archbishop of valencia",
    "gioacchino torriani,",
    "rené de prie, cardinal priest of santa lucia in septisolio",
    "giovanni conti, cardinal deacon of s. nereo e achilleo",
    "“father dioneo farnese,” aka. “la bella” julia farnese",
    "pierre d’aubusson, crusader general",
    "captain micheletto corella of the papal guard",
    "francesco alidosi",
    "bocca (della faggiuolo) “il pisano”",
    "caterina sforza riario, lady of forlì",
    "cardinal giovanni michiel zen",
    "paolo campofregoso, cardinal deacon of sant’ anastasia",
    "agostino tillio",
    "adriano “hadrian” castellesi",
    "rodrigo borgia, cardinal deacon of san nicola in carcere",
    "brother vincenzo bandello",
    "jorge da costa, cardinal of portugal",
    "antonio pallavicini gentili",
    "anne, duchess of brittany",
    "ippolito d’este, cardinal of hungary and ferrara",
    "raffaele riario (della rovere)",
    "franceschetto cybo, count of the lateran palace",
    "girolamo basso della rovere, cardinal priest of san crisogono",
    "francesco “samson” nanni, minister general of the franciscans",
    "beatrice of naples, twice queen of hungary, bohemia, luxemburg, lusatia, moravia, austria, galicia, silesia, dalmatia, and besterce, of the house of trastámara.",
    "giovanni colonna, cardinal deacon of s. maria in aquiro",
    "louis of orleans, king of france and naples, and duke of milan",
    "doctor john argentine",
    "giuliano della rovere, cardinal deacon of san pietro in vincula",
    "giovanni battista zen",
    "louis of orleans, regent of france and naples, duke of orleans and milan",
    "galahad sanseverino",
    "johann burchard",
    "ascanio maria visconti sforza, cardinal of cremona",
    "cardinal oliverio carafa,",
    "guillaume briçonnet, cardinal of st. malo",
    "brother constantine, aka. prince cem",
    "matthäus schiner, cardinal of switzerland",
    "pietro bembo, cardinal priest of san clemente",
    "maximilian i, of the house of hapsburg, by birth archduke of austria, by marriage regent-protector of burgundy, and by election emperor, king of the germans, king of lombardy, and king of the romans.",
    "francesco todeschini piccolomini, cardinal of siena",
    "maffeo gherardi, patriarch of venice",
    "michel de bucy, cardinal of bourges",
    "giovanni “leo” de medici, cardinal of florence",
    "brother tillio, apostolic secretary",
    "josquin lebloitte dit des prez",
    "giovanni battista savelli, cardinal deacon of ss vito e modesto",
    "gherardo maffei",
    "francesco soderini, cardinal bishop of sabina",
    "henry vii, king of england and of france, lord of ireland, of the house of tudor",
    "giovanni “achilles” sangiorgio (sometimes “sangiorgi”)",
    "adriano “hadrian” castellesi",
    "brother annio of viterbo (vote-counter)",
    "raffaello maffei “volterrano”",
    # "ardicino della porta junior",
    "michele tarchaniota marullo",
    "the roman mob"
]

def testing():
    return 1

def parser():
    terms = ["close ally", "ally", "allies", "confidant", "client", "correspondent", "patron", "boss", "colleague"]
    characters = {}
    directory = "Character_Sheets_ALL_2024/"
    for file_name in os.listdir(directory):
        document = Document(directory + file_name)
        extract = document.paragraphs
        name = (extract[0].text.split("From")[0][6:].rstrip())
        if "ardicino della porta" in name.lower():
            continue
        current = Character(name)
        current.other = extract[0].text
        for table in document.tables:
            for row in table.rows[1:]:
                data = row.cells
                if data[0].text != data[1].text:
                    category = ""
                    relation = data[1].text.lower().replace("allies", "ally")
                    for term in terms:
                        if term in relation:
                            category = term
                            break
                    if "close ally" in relation:
                        category = "close ally"
                    if "ex-patron" in relation:
                        category = "ally"
                    if "ally (boss)" in relation or "ally (former boss)" in relation:
                        category = "ally"
                    if category != "":
                        if category in current.relations:
                            current.relations[category].append(data[0].text)
                        else:
                            current.relations[category] = [(data[0].text)]
                            current.relationsFiltered[category] = []
                    if "are not vatican employees so are not your colleagues" in data[0].text.lower():
                        print(data[1].text)
        characters[name.lower()] = current

    characters["the roman mob"] = Character("The Roman Mob")
    count = 0
    other = 0
    filterone = [" of ", " the ", " in ", " and ", " de ", "da ", " a ", " della "]
    filtertwo = [
                    ["cpt.", "captain"],
                    ["capta", "captain"],
                    ["bocca", "bocca (della faggiuolo)"],
                    ["cardal", "cardinal"],
                    ["alessro", "alessandro"], 
                    ["josqu", "josquin lebloitte"], 
                    ["piccolomi", "francesco todeschini piccolomini"], 
                    ["bror", "brother"], 
                    ["valentino", "cesare “valentino” borgia, archbishop"], 
                    ["marullo", "michele tarchaniota marullo"],
                    ["brother tillio", "brother tillio, apostolic secretary"], 
                    ["far alidosi", "francesco"], 
                    ["anne  brittany", "anne, duchess of brittany"], 
                    ["cardinals", "cardinal"], 
                    ["cardinal orsi", "cardinal orsini"], 
                    ["cardinal sansevero", "cardinal sanseverino"], 
                    ["corella", "captain corella"], 
                    ["farnese", "“father dioneo farnese"],
                    ["cardinal orsinini (grrr….)", "giambattista orsini, cardinal deacon"],
                    ["cardinal sanseverino (nephew giovanni)", "cardinal sanseverino"],
                    ["sanseverino", "federico sanseverino"],
                    ["porta", "ardicino porta junior"],
                    ["constantine", "brother constantine prince cem"],
                    ["alidosi", "francesco alidosi"],
                    ["argentine", "doctor john argentine"],
                    ["orsinini", "giambattista orsini"],
                    ["soderini", "francesco soderini"],
                    ["torriani", "gioacchino torriani inquisitor general"],
                    ["client charles", "charles viii"],
                    ["client kinsman charles", "charles viii"],
                    ["gentili", "antonio pallavicini gentili"],
                    ["tillio", "tillio, apostolic secretary"],
                    ["johan burchard", "johann burchard"],
                    ["josquin lebloittein", "josquin lebloitte dit des prez"],
                    ["conti", "giovanni conti achilleo"],
                    ["raffaelle riario", "raffaele riario"],
                    ["sangiorgio", "giovanni “achilles” sangiorgi"],
                    ["beatrice", "beatrice naples hungary"],
                    ["agentine", "doctor john argentine"],
                    ["vinovo", "domenico di vinovo"],
                    ["torianni", "gioacchino torriani inquisitor general"],
                    ["raffaello riario", "raffaele riario"],
                    ["alessandro", "alessandro farnese"],
                    ["isabella", "isabella i castile"],
                    ["riario)", "raffaele riario"],
                    ["domenico", "domenico vinovo"],
                    ["raffaello", "raffaello maffei “volterrano”"],
                    ["mario", "father mario maffei"],
                    ["gherardo", "gherardo maffei"],
                    ["patriarch", "maffeo gherardi patriarch"],
                    ["campofregoso", "paolo campofregoso anastasia"],
                    ["burchard", "johann burchard"],
                ]
    for character in characters:
        for key in characters[character].relations:
            for itm in characters[character].relations[key]:
                match = ""
                naming = itm.lower()
                naming = naming.replace(" & ", ", ")
                naming = naming.split(',')
                for name in naming:
                    limit = 1
                    if len(name) > 2:
                        flag = True
                        for itm in filterone:
                            if len(name.split()) > 2:
                                name = name.replace(itm, " ")
                        for itm in filtertwo:
                            name = name.replace(itm[0], itm[1])
                        while name[0] == " " and len(name) > 2:
                            name = name[1:]
                        for person in people:
                            match_count = len(set(re.findall(r'\b\w+\b', name)) & set(re.findall(r'\b\w+\b', person.replace(',', ''))))
                            if match_count > limit and "non-player-character" not in name:
                                flag = False
                                characters[character].relationsFiltered[key].append(characters[person])
                                limit = 2
                        if flag:
                            if "crowned heads" in name or "monarchs" in name:
                                for a in people:
                                    if "king" in a or "queen" in a or "prince" in a:
                                        if a != character.lower():
                                            characters[character].relationsFiltered[key].append(characters[a])
                            elif "all frenchmen" in name:
                                for a in people:
                                    if "france" in characters[a].other.lower():
                                        if a != character.lower():
                                            characters[character].relationsFiltered[key].append(characters[a])
                            elif "the maffeis" in name:
                                    for a in people:
                                        if "maffei" in a:
                                            if a != character.lower():
                                                characters[character].relationsFiltered[key].append(characters[a])
                            elif "literally everyone" in name:
                                for a in people:
                                    if a != character.lower():
                                        characters[character].relationsFiltered[key].append(characters[a])

            # if characters["ardicino della porta junior"] not in characters[character].relationsFiltered["ally"]:
            #     characters[character].relationsFiltered["ally"].append(characters["ardicino della porta junior"])

    return characters
